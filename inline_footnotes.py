#!/usr/bin/env python3
"""
Inline Footnotes and Endnotes
Processes Word documents to inline footnotes and endnotes content at their reference points.
This ensures that scripture references and other important information in notes are 
available for indexing and processing.
"""

import os
import sys
import argparse
import shutil
import re
from pathlib import Path
from docx import Document

def extract_manual_notes(doc: Document) -> dict:
    """Extract manually inserted notes from the end of a document"""
    notes = {}
    
    # Look through all paragraphs, starting from the end
    for para in reversed(doc.paragraphs):
        text = para.text.strip()
        
        if not text:  # Skip empty paragraphs
            continue
            
        # Look for note patterns like "[1] content", "[2] more content", etc.
        match = re.match(r'^\[(\d+)\]\s*(.*)', text)
        if match:
            note_num = match.group(1)
            note_content = match.group(2).strip()
            
            if note_num and note_content:
                notes[note_num] = note_content
                print(f"Found note {note_num}: {note_content[:50]}...")
        else:
            # If we encounter a paragraph that doesn't match the note pattern
            # and we already have some notes, we're probably done with the notes section
            if notes:
                break
                
    return notes

def inline_manual_notes(input_path: str, output_path: str = None):
    """Inline manually created footnotes/endnotes"""
    
    if output_path is None:
        path_obj = Path(input_path)
        output_path = str(path_obj.parent / f"{path_obj.stem}_inlined{path_obj.suffix}")
    
    # Backup original
    backup_path = input_path + ".backup"
    if not os.path.exists(backup_path):
        shutil.copy2(input_path, backup_path)
        print(f"Created backup: {backup_path}")
    
    try:
        doc = Document(input_path)
        
        # Extract manual notes
        notes = extract_manual_notes(doc)
        
        if not notes:
            print("No manual notes found. Copying file as-is.")
            shutil.copy2(input_path, output_path)
            return True
        
        print(f"Found {len(notes)} manual notes to inline")
        
        # Process paragraphs to inline notes
        for para in doc.paragraphs:
            if para.text.strip():
                # Use finditer to handle multiple references in one paragraph
                for note_num, note_content in notes.items():
                    # Look for bracketed references like [1], [2]
                    pattern = r'\[' + re.escape(note_num) + r'\]'
                    
                    # Replace all occurrences
                    if re.search(pattern, para.text):
                        print(f"  Inlining note {note_num}...")
                        para.text = re.sub(pattern, f" ({note_content}) ", para.text)
                        
        # Remove the original notes section at the end of the document
        notes_section_found = False
        for para in reversed(doc.paragraphs):
            if re.match(r'^(notes|endnotes|footnotes)\b', para.text.strip(), re.IGNORECASE):
                notes_section_found = True
            
            if notes_section_found:
                # Remove the note paragraphs
                p_element = para._p
                p_element.getparent().remove(p_element)
        
        # Save the processed document
        doc.save(output_path)
        print(f"Processed document saved to: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error processing document: {e}")
        return False

def main():
    parser = argparse.ArgumentParser(description='Inline manual footnotes/endnotes in Word documents')
    parser.add_argument('input', help='Input Word document')
    parser.add_argument('-o', '--output', help='Output file path')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.input):
        print(f"Input file not found: {args.input}")
        sys.exit(1)
    
    success = inline_manual_notes(args.input, args.output)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
