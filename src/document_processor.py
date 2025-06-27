# Enhanced Document Processing with Image and Table Support

import os
import re
import logging
import pandas as pd
import PyPDF2
import docx
import pdfplumber
import pytesseract
from PIL import Image
import tabula
import magic
import io
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm
from .database_manager import DatabaseManager
from .document_preprocessor import DocumentPreprocessor

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = ['.txt', '.pdf', '.doc', '.docx', '.csv', '.jpg', '.jpeg', '.png', '.tiff', '.bmp']

class DocumentProcessor:
    def __init__(self, config: Dict[str, any], db_manager: DatabaseManager):
        self.config = config
        self.db_manager = db_manager
        self.max_chunk_size = config['database']['max_chunk_size']
        self.chunk_overlap = config['database']['chunk_overlap']
        
        # Initialize document preprocessor
        self.preprocessor = DocumentPreprocessor()
        logger.info("Initialized document preprocessor for text normalization")

    def process_directory(self, directory: str):
        """Process all documents in a directory."""
        if not os.path.exists(directory):
            logger.warning(f"Directory {directory} does not exist")
            return
            
        processed_count = 0
        skipped_count = 0
        error_count = 0
        
        for root, _, files in os.walk(directory):
            for file in files:
                file_path = os.path.join(root, file)
                if self.is_supported(file_path):
                    # Check if document is already processed
                    if self.db_manager.is_document_already_processed(file_path):
                        logger.info(f"Skipping already processed file: {file}")
                        skipped_count += 1
                        continue
                        
                    logger.info(f"Processing file: {file}")
                    try:
                        doc_text = self.extract_text(file_path)
                        if not doc_text.strip():
                            logger.warning(f"No text extracted from {file_path}")
                            continue
                            
                        chunks = self.chunk_text(doc_text)
                        if not chunks:
                            logger.warning(f"No chunks created from {file_path}")
                            continue
                            
                        result = self.db_manager.add_document(file_path, file, self.get_file_type(file_path), chunks)
                        if result:
                            processed_count += 1
                            logger.info(f"Successfully processed file: {file}")
                        else:
                            error_count += 1
                            logger.error(f"Failed to add file: {file} to database")
                            
                    except Exception as e:
                        error_count += 1
                        logger.error(f"Failed to process file: {file}: {e}")
                        
        logger.info(f"Processing complete: {processed_count} files processed, {skipped_count} files skipped, {error_count} errors")

    def process_specific_files(self, file_paths: List[str], force: bool = False):
        """Process a specific list of files."""
        processed_count = 0
        skipped_count = 0
        error_count = 0
        
        for file_path in file_paths:
            if not os.path.exists(file_path):
                logger.warning(f"File {file_path} does not exist")
                error_count += 1
                continue
                
            if not self.is_supported(file_path):
                logger.warning(f"File format not supported: {file_path}")
                error_count += 1
                continue
                
            # Check if document is already processed (unless force is True)
            if not force and self.db_manager.is_document_already_processed(file_path):
                logger.info(f"Skipping already processed file: {os.path.basename(file_path)}")
                skipped_count += 1
                continue
                
            logger.info(f"Processing file: {os.path.basename(file_path)}")
            try:
                doc_text = self.extract_text(file_path)
                if not doc_text.strip():
                    logger.warning(f"No text extracted from {file_path}")
                    error_count += 1
                    continue
                    
                chunks = self.chunk_text(doc_text)
                if not chunks:
                    logger.warning(f"No chunks created from {file_path}")
                    error_count += 1
                    continue
                    
                result = self.db_manager.add_document(file_path, os.path.basename(file_path), self.get_file_type(file_path), chunks)
                if result:
                    processed_count += 1
                    logger.info(f"Successfully processed file: {os.path.basename(file_path)} ({len(chunks)} chunks)")
                else:
                    error_count += 1
                    logger.error(f"Failed to add file: {os.path.basename(file_path)} to database")
                    
            except Exception as e:
                error_count += 1
                logger.error(f"Failed to process file: {os.path.basename(file_path)}: {e}")
                
        logger.info(f"Specific file processing complete: {processed_count} files processed, {skipped_count} files skipped, {error_count} errors")
        return {"processed": processed_count, "skipped": skipped_count, "errors": error_count}

    def is_supported(self, file_path: str) -> bool:
        """Check if the file format is supported."""
        return any(file_path.endswith(ext) for ext in SUPPORTED_FORMATS)

    def get_file_type(self, file_path: str) -> str:
        """Detect the MIME type of a file."""
        mime = magic.Magic(mime=True)
        return mime.from_file(file_path)

    def extract_text(self, file_path: str) -> str:
        """Extract text, tables, and images from different document formats."""
        raw_text = ""

        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                raw_text = f.read()
        elif file_path.endswith('.pdf'):
            raw_text = self.extract_pdf(file_path)
        elif file_path.endswith(('.doc', '.docx')):
            raw_text = self.extract_word(file_path)
        elif file_path.endswith('.csv'):
            raw_text = self.extract_csv(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path}")

        # Apply preprocessing to normalize scripture references and theological terms
        try:
            preprocessed_text = self.preprocessor.preprocess_document(raw_text)
            logger.debug(f"Applied preprocessing to {file_path}")
            return preprocessed_text
        except Exception as e:
            logger.warning(f"Preprocessing failed for {file_path}: {e}. Using raw text.")
            return raw_text

    def extract_pdf(self, file_path: str) -> str:
        """Extract text and tables from a PDF file."""
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
                tables = page.extract_tables()
                for table in tables:
                    text += "\n" + self.format_table(table)
        return text

    def extract_word(self, file_path: str) -> str:
        """Extract text and tables from a Word document."""
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        for table in doc.tables:
            text += "\n" + self.format_table([[cell.text for cell in row.cells] for row in table.rows])
        return text

    def extract_csv(self, file_path: str) -> str:
        """Extract text from a CSV file."""
        df = pd.read_csv(file_path)
        return df.to_string()

    def format_table(self, table: list) -> str:
        """Format table data as text."""
        formatted_table = ""
        for row in table:
            # Handle None values in table cells
            cleaned_row = [str(cell) if cell is not None else "" for cell in row]
            formatted_table += "\t".join(cleaned_row) + "\n"
        return formatted_table

    def chunk_text(self, text: str) -> List[str]:
        """Chunk text into smaller pieces for storage."""
        words = text.split()
        chunks = []
        for start in range(0, len(words), self.max_chunk_size - self.chunk_overlap):
            end = start + self.max_chunk_size
            chunk = " ".join(words[start:end])
            chunks.append(chunk)
        return chunks

    def extract_image_text(self, file_path: str) -> str:
        """Extract text from images using OCR."""
        text = ""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            logger.info(f"Extracted text from image in {file_path}")
        except Exception as e:
            logger.error(f"Failed to extract image text from {file_path}: {e}")
        return text
