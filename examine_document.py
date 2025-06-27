#!/usr/bin/env python3
"""
Examine Document Structure
Simple script to examine the structure of a Word document and identify footnotes/endnotes
"""

import sys
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

def examine_document(doc_path):
    print(f"Examining: {doc_path}")
    print("=" * 50)
    
    # Load document
    doc = Document(doc_path)
    
    # Show basic info
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    
    # Look for patterns that might indicate footnotes
    footnote_indicators = []
    superscript_found = False
    
    for i, para in enumerate(doc.paragraphs[:20]):  # First 20 paragraphs
        text = para.text.strip()
        if text:
            print(f"\nParagraph {i+1}: {text[:100]}...")
            
            # Check each run for superscript or note indicators
            for run in para.runs:
                if run.font.superscript:
                    superscript_found = True
                    print(f"  -> Superscript text found: '{run.text}'")
                
                # Look for common footnote patterns
                import re
                patterns = [
                    r'\d+',  # Numbers
                    r'[¹²³⁴⁵⁶⁷⁸⁹⁰]',  # Superscript numbers
                    r'\[\d+\]',  # [1], [2], etc.
                    r'\(\d+\)',  # (1), (2), etc.
                ]
                
                for pattern in patterns:
                    matches = re.findall(pattern, run.text)
                    if matches:
                        footnote_indicators.extend(matches)
    
    print(f"\nSuperscript formatting found: {superscript_found}")
    print(f"Potential footnote indicators: {set(footnote_indicators)}")
    
    # Try to examine the document's XML structure
    print("\n" + "=" * 50)
    print("EXAMINING XML STRUCTURE")
    print("=" * 50)
    
    try:
        # Open as ZIP file to examine structure
        with zipfile.ZipFile(doc_path, 'r') as zip_file:
            files = zip_file.namelist()
            print("Files in document:")
            for file in files:
                print(f"  {file}")
            
            # Check for footnotes.xml
            if 'word/footnotes.xml' in files:
                print("\nFound footnotes.xml!")
                footnotes_content = zip_file.read('word/footnotes.xml')
                print("Footnotes content preview:")
                print(footnotes_content[:500].decode('utf-8', errors='ignore'))
            
            # Check for endnotes.xml
            if 'word/endnotes.xml' in files:
                print("\nFound endnotes.xml!")
                endnotes_content = zip_file.read('word/endnotes.xml')
                print("Endnotes content preview:")
                print(endnotes_content[:500].decode('utf-8', errors='ignore'))
            
            # Examine document.xml for footnote references
            if 'word/document.xml' in files:
                print("\nExamining document.xml for footnote references...")
                doc_content = zip_file.read('word/document.xml')
                doc_xml = doc_content.decode('utf-8', errors='ignore')
                
                # Look for footnote references
                if 'footnoteReference' in doc_xml:
                    print("Found footnoteReference elements in document!")
                if 'endnoteReference' in doc_xml:
                    print("Found endnoteReference elements in document!")
                
                # Count occurrences
                footnote_refs = doc_xml.count('footnoteReference')
                endnote_refs = doc_xml.count('endnoteReference')
                print(f"Footnote references: {footnote_refs}")
                print(f"Endnote references: {endnote_refs}")
                
    except Exception as e:
        print(f"Error examining XML structure: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python3 examine_document.py <document_path>")
        sys.exit(1)
    
    examine_document(sys.argv[1])
